#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para gerar manuais de usuário da plataforma AISAM
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Cores da AISAM (convertidas de HSL para RGB)
AZUL_PRINCIPAL = RGBColor(0, 102, 230)  # #0066E6
AZUL_ESCURO = RGBColor(0, 82, 184)      # #0052B8
AZUL_CLARO = RGBColor(51, 153, 255)     # #3399FF
CINZA_TEXTO = RGBColor(51, 51, 51)


def add_page_break(doc):
    """Adiciona quebra de página"""
    doc.add_page_break()


def add_cover_page(doc, title, subtitle, logo_path=None):
    """Adiciona capa profissional com logo"""
    # Espaçamento superior
    for _ in range(5):
        doc.add_paragraph()

    # Logo (se existir)
    if logo_path and os.path.exists(logo_path):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_logo = p_logo.add_run()
        try:
            run_logo.add_picture(logo_path, width=Inches(3.0))
        except:
            # Se falhar ao carregar o logo, continua sem ele
            pass
        doc.add_paragraph()
        doc.add_paragraph()

    # Título principal
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(title)
    run_title.font.size = Pt(32)
    run_title.font.bold = True
    run_title.font.color.rgb = AZUL_PRINCIPAL

    doc.add_paragraph()

    # Subtítulo
    p_subtitle = doc.add_paragraph()
    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_subtitle = p_subtitle.add_run(subtitle)
    run_subtitle.font.size = Pt(18)
    run_subtitle.font.color.rgb = AZUL_ESCURO

    # Espaçamento
    for _ in range(8):
        doc.add_paragraph()

    # Informações adicionais
    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_info = p_info.add_run('Plataforma de Gestão de Vagas e Currículos')
    run_info.font.size = Pt(14)
    run_info.font.color.rgb = CINZA_TEXTO

    doc.add_paragraph()

    p_version = doc.add_paragraph()
    p_version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_version = p_version.add_run('Versão 1.0 - 2026')
    run_version.font.size = Pt(12)
    run_version.font.color.rgb = CINZA_TEXTO

    # Espaçamento inferior
    for _ in range(3):
        doc.add_paragraph()

    # Rodapé da capa
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_footer = p_footer.add_run('AISAM - Associação das Indústrias de Santa Maria da Serra e Região')
    run_footer.font.size = Pt(10)
    run_footer.font.color.rgb = CINZA_TEXTO
    run_footer.font.italic = True


def add_logo_header(doc, title):
    """Adiciona logo e título ao documento"""
    # Título principal
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.color.rgb = AZUL_PRINCIPAL
        run.font.size = Pt(24)
        run.font.bold = True

    # Linha separadora
    doc.add_paragraph('_' * 80)
    doc.add_paragraph()


def add_image_placeholder(doc, description, height=3.0):
    """Adiciona espaço para inserção de imagem com descrição"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Caixa para imagem
    run = p.add_run()
    run.add_text('┌' + '─' * 78 + '┐\n')
    run.add_text('│' + ' ' * 78 + '│\n')
    run.add_text('│' + ' ' * 25 + '[INSERIR PRINT DE TELA AQUI]' + ' ' * 24 + '│\n')
    run.add_text('│' + ' ' * 78 + '│\n')
    run.add_text('└' + '─' * 78 + '┘')
    run.font.color.rgb = AZUL_CLARO
    run.font.size = Pt(9)

    # Descrição da imagem
    desc = doc.add_paragraph(description)
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in desc.runs:
        run.font.color.rgb = CINZA_TEXTO
        run.font.size = Pt(10)
        run.font.italic = True

    doc.add_paragraph()


def add_section_title(doc, title, level=1):
    """Adiciona título de seção"""
    heading = doc.add_heading(title, level)
    for run in heading.runs:
        run.font.color.rgb = AZUL_PRINCIPAL if level == 1 else AZUL_ESCURO
        run.font.bold = True


def add_styled_paragraph(doc, text, bold=False, color=None):
    """Adiciona parágrafo com estilo"""
    p = doc.add_paragraph(text)
    for run in p.runs:
        if bold:
            run.font.bold = True
        if color:
            run.font.color.rgb = color
        run.font.size = Pt(11)
    return p


def add_bullet_list(doc, items):
    """Adiciona lista com bullets"""
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(11)


def add_numbered_list(doc, items):
    """Adiciona lista numerada"""
    for item in items:
        p = doc.add_paragraph(item, style='List Number')
        for run in p.runs:
            run.font.size = Pt(11)


def add_table_simple(doc, headers, rows):
    """Adiciona tabela simples"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Cabeçalhos
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = AZUL_PRINCIPAL

    # Dados
    for i, row in enumerate(rows):
        row_cells = table.rows[i + 1].cells
        for j, cell_text in enumerate(row):
            row_cells[j].text = cell_text


def criar_manual_admin():
    """Cria manual do Administrador AISAM"""
    doc = Document()

    # Configurações do documento
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Capa com logo
    logo_path = 'frontend/src/assets/aisam-logo.webp'
    add_cover_page(doc, 'MANUAL DO ADMINISTRADOR', 'Sistema de Gestão AISAM', logo_path)

    add_page_break(doc)

    # Página de apresentação
    add_logo_header(doc, 'MANUAL DO USUÁRIO\nADMINISTRADOR AISAM')

    add_styled_paragraph(doc, 'Plataforma de Gestão de Vagas e Currículos', bold=True)
    add_styled_paragraph(doc, 'Versão 1.0 - 2026')
    add_styled_paragraph(doc, '')
    add_styled_paragraph(doc, 'Este manual destina-se aos Administradores da plataforma AISAM.')

    add_page_break(doc)

    # Índice
    add_section_title(doc, 'ÍNDICE', 1)
    add_bullet_list(doc, [
        '1. Introdução',
        '2. Acesso ao Sistema',
        '3. Dashboard Administrativo',
        '4. Gestão de Recrutadores',
        '5. Gestão de Associados (Empresas)',
        '6. Gestão de Vagas',
        '7. Gestão de Áreas de Atuação',
        '8. Auditoria e Segurança',
        '9. Relatórios',
        '10. Perguntas Frequentes'
    ])

    add_page_break(doc)

    # 1. Introdução
    add_section_title(doc, '1. INTRODUÇÃO', 1)
    add_styled_paragraph(doc, 'Bem-vindo ao sistema de gestão da AISAM!')
    add_styled_paragraph(doc, 'Como Administrador, você tem acesso completo a todas as funcionalidades da plataforma, incluindo:')
    add_bullet_list(doc, [
        'Gerenciamento completo de recrutadores e associados',
        'Controle total sobre vagas publicadas',
        'Configuração de áreas de atuação do sistema',
        'Acesso a logs de auditoria e segurança',
        'Geração de relatórios gerenciais'
    ])

    add_section_title(doc, '1.1 Sobre este Manual', 2)
    add_styled_paragraph(doc, 'Este documento fornece instruções detalhadas sobre todas as funcionalidades disponíveis para o perfil de Administrador.')

    add_page_break(doc)

    # 2. Acesso ao Sistema
    add_section_title(doc, '2. ACESSO AO SISTEMA', 1)

    add_section_title(doc, '2.1 Fazendo Login', 2)
    add_styled_paragraph(doc, 'Para acessar o sistema como Administrador:')
    add_numbered_list(doc, [
        'Acesse a URL da plataforma AISAM',
        'Clique em "Acessar Sistema"',
        'Selecione "Login Administrador"',
        'Digite seu e-mail e senha cadastrados',
        'Clique em "Entrar"'
    ])

    add_image_placeholder(doc, 'Figura 1: Tela de login do Administrador')

    add_section_title(doc, '2.2 Recuperação de Senha', 2)
    add_styled_paragraph(doc, 'Caso tenha esquecido sua senha:')
    add_numbered_list(doc, [
        'Na tela de login, clique em "Esqueci minha senha"',
        'Digite seu e-mail cadastrado',
        'Você receberá um e-mail com link para redefinir sua senha',
        'Clique no link e defina uma nova senha',
        'Faça login com a nova senha'
    ])

    add_image_placeholder(doc, 'Figura 2: Tela de recuperação de senha')

    add_page_break(doc)

    # 3. Dashboard
    add_section_title(doc, '3. DASHBOARD ADMINISTRATIVO', 1)
    add_styled_paragraph(doc, 'Após fazer login, você será direcionado ao Dashboard Administrativo, que oferece acesso rápido a todas as funcionalidades principais.')

    add_image_placeholder(doc, 'Figura 3: Dashboard do Administrador', height=4.0)

    add_section_title(doc, '3.1 Cards de Acesso Rápido', 2)
    add_styled_paragraph(doc, 'O dashboard contém os seguintes cards:')
    add_bullet_list(doc, [
        'Gerenciamento de Usuários (Recrutadores)',
        'Gerenciamento de Associados (Empresas)',
        'Gerenciamento de Vagas',
        'Áreas de Atuação',
        'Banco de Currículos',
        'Auditoria do Sistema',
        'Relatórios Gerenciais'
    ])

    add_page_break(doc)

    # 4. Gestão de Recrutadores
    add_section_title(doc, '4. GESTÃO DE RECRUTADORES', 1)
    add_styled_paragraph(doc, 'Nesta seção você pode gerenciar todos os recrutadores cadastrados na plataforma.')

    add_section_title(doc, '4.1 Visualizar Lista de Recrutadores', 2)
    add_numbered_list(doc, [
        'No dashboard, clique em "Gerenciamento de Usuários"',
        'Você verá a lista completa de recrutadores cadastrados',
        'Use os filtros para buscar por nome, email ou status'
    ])

    add_image_placeholder(doc, 'Figura 4: Lista de recrutadores cadastrados')

    add_section_title(doc, '4.2 Convidar Novo Recrutador', 2)
    add_styled_paragraph(doc, 'Para adicionar um novo recrutador ao sistema:')
    add_numbered_list(doc, [
        'Clique no botão "Convidar Recrutador"',
        'Preencha os dados do recrutador:',
        '  - Nome completo',
        '  - E-mail profissional',
        '  - Associado vinculado (opcional)',
        'Clique em "Enviar Convite"',
        'O recrutador receberá um e-mail com link para ativar sua conta',
        'O link de convite tem validade de 7 dias'
    ])

    add_image_placeholder(doc, 'Figura 5: Formulário de convite para novo recrutador')

    add_section_title(doc, '4.3 Ativar/Desativar Recrutador', 2)
    add_styled_paragraph(doc, 'Para gerenciar o status de um recrutador:')
    add_numbered_list(doc, [
        'Na lista de recrutadores, localize o usuário desejado',
        'Clique no botão de ações (três pontos)',
        'Selecione "Ativar" ou "Desativar"',
        'Confirme a ação'
    ])
    add_styled_paragraph(doc, 'Nota: Recrutadores desativados não conseguem fazer login no sistema.', bold=True, color=AZUL_ESCURO)

    add_image_placeholder(doc, 'Figura 6: Opções de ativação/desativação de recrutador')

    add_section_title(doc, '4.4 Editar Dados do Recrutador', 2)
    add_numbered_list(doc, [
        'Clique sobre o recrutador na lista',
        'Visualize os detalhes completos',
        'Clique em "Editar"',
        'Atualize os campos necessários',
        'Clique em "Salvar"'
    ])

    add_image_placeholder(doc, 'Figura 7: Edição de dados do recrutador')

    add_page_break(doc)

    # 5. Gestão de Associados
    add_section_title(doc, '5. GESTÃO DE ASSOCIADOS (EMPRESAS)', 1)
    add_styled_paragraph(doc, 'Gerencie as empresas associadas à AISAM.')

    add_section_title(doc, '5.1 Visualizar Associados', 2)
    add_numbered_list(doc, [
        'No dashboard, clique em "Gerenciamento de Associados"',
        'Visualize a lista de todas as empresas associadas',
        'Use filtros para buscar por nome, CNPJ ou status'
    ])

    add_image_placeholder(doc, 'Figura 8: Lista de associados cadastrados')

    add_section_title(doc, '5.2 Cadastrar Novo Associado', 2)
    add_numbered_list(doc, [
        'Clique em "Novo Associado"',
        'Preencha o formulário com os dados da empresa:',
        '  - Razão Social',
        '  - Nome Fantasia',
        '  - CNPJ',
        '  - E-mail',
        '  - Telefone',
        '  - Endereço completo (Rua, Cidade, Estado, CEP)',
        'Clique em "Cadastrar"'
    ])

    add_image_placeholder(doc, 'Figura 9: Formulário de cadastro de associado')

    add_section_title(doc, '5.3 Aprovar Pedidos de Associação', 2)
    add_styled_paragraph(doc, 'Empresas podem solicitar associação através do site. Para aprovar:')
    add_numbered_list(doc, [
        'Acesse "Pedidos Pendentes"',
        'Revise os dados da empresa solicitante',
        'Clique em "Aprovar" ou "Rejeitar"',
        'Adicione observações se necessário',
        'Confirme a ação'
    ])

    add_image_placeholder(doc, 'Figura 10: Aprovação de pedido de associação')

    add_section_title(doc, '5.4 Editar Dados do Associado', 2)
    add_numbered_list(doc, [
        'Selecione o associado na lista',
        'Clique em "Editar"',
        'Atualize as informações necessárias',
        'Clique em "Salvar"'
    ])

    add_image_placeholder(doc, 'Figura 11: Edição de dados do associado')

    add_page_break(doc)

    # 6. Gestão de Vagas
    add_section_title(doc, '6. GESTÃO DE VAGAS', 1)
    add_styled_paragraph(doc, 'Visualize e gerencie todas as vagas publicadas na plataforma.')

    add_section_title(doc, '6.1 Visualizar Todas as Vagas', 2)
    add_numbered_list(doc, [
        'No dashboard, clique em "Gerenciamento de Vagas"',
        'Visualize todas as vagas publicadas por recrutadores',
        'Use filtros para buscar por:',
        '  - Status (Aberta, Pausada, Fechada, Arquivada)',
        '  - Associado',
        '  - Área de Atuação',
        '  - Regime de Trabalho',
        '  - Senioridade'
    ])

    add_image_placeholder(doc, 'Figura 12: Lista completa de vagas')

    add_section_title(doc, '6.2 Detalhes da Vaga', 2)
    add_numbered_list(doc, [
        'Clique sobre uma vaga para ver detalhes completos',
        'Visualize:',
        '  - Descrição completa',
        '  - Recrutador responsável',
        '  - Associado vinculado',
        '  - Número de candidaturas recebidas',
        '  - Status atual',
        '  - Datas de criação e atualização'
    ])

    add_image_placeholder(doc, 'Figura 13: Detalhes de uma vaga')

    add_section_title(doc, '6.3 Moderar Vagas', 2)
    add_styled_paragraph(doc, 'Como administrador, você pode moderar vagas se necessário:')
    add_bullet_list(doc, [
        'Pausar vagas inadequadas',
        'Encerrar vagas vencidas',
        'Editar informações incorretas',
        'Arquivar vagas antigas'
    ])

    add_image_placeholder(doc, 'Figura 14: Opções de moderação de vagas')

    add_page_break(doc)

    # 7. Áreas de Atuação
    add_section_title(doc, '7. GESTÃO DE ÁREAS DE ATUAÇÃO', 1)
    add_styled_paragraph(doc, 'Configure as áreas de atuação disponíveis para categorização de vagas e currículos.')

    add_section_title(doc, '7.1 Visualizar Áreas', 2)
    add_numbered_list(doc, [
        'No dashboard, clique em "Áreas de Atuação"',
        'Visualize todas as áreas cadastradas',
        'Veja quantas vagas e candidatos estão vinculados a cada área'
    ])

    add_image_placeholder(doc, 'Figura 15: Lista de áreas de atuação')

    add_section_title(doc, '7.2 Criar Nova Área', 2)
    add_numbered_list(doc, [
        'Clique em "Nova Área de Atuação"',
        'Digite o nome da área (ex: Tecnologia da Informação, Logística, Recursos Humanos)',
        'Adicione uma descrição (opcional)',
        'Clique em "Criar"'
    ])

    add_image_placeholder(doc, 'Figura 16: Formulário de criação de área de atuação')

    add_section_title(doc, '7.3 Editar ou Excluir Área', 2)
    add_numbered_list(doc, [
        'Clique sobre a área desejada',
        'Para editar: altere o nome/descrição e clique em "Salvar"',
        'Para excluir: clique em "Excluir" e confirme'
    ])
    add_styled_paragraph(doc, 'Atenção: Não é possível excluir áreas que possuem vagas ou candidatos vinculados.', bold=True, color=AZUL_ESCURO)

    add_image_placeholder(doc, 'Figura 17: Edição de área de atuação')

    add_page_break(doc)

    # 8. Auditoria
    add_section_title(doc, '8. AUDITORIA E SEGURANÇA', 1)
    add_styled_paragraph(doc, 'Monitore todas as ações realizadas no sistema através dos logs de auditoria.')

    add_section_title(doc, '8.1 Acessar Logs de Auditoria', 2)
    add_numbered_list(doc, [
        'No dashboard, clique em "Auditoria do Sistema"',
        'Visualize o registro cronológico de todas as ações',
        'Use filtros para buscar por:',
        '  - Tipo de ação (CREATE, UPDATE, DELETE, LOGIN, LOGOUT, etc.)',
        '  - Usuário',
        '  - Data/período',
        '  - Tipo de entidade (Vaga, Candidato, Recrutador, etc.)'
    ])

    add_image_placeholder(doc, 'Figura 18: Logs de auditoria do sistema')

    add_section_title(doc, '8.2 Tipos de Ações Registradas', 2)
    add_styled_paragraph(doc, 'O sistema registra as seguintes ações:')

    add_table_simple(doc,
        ['Ação', 'Descrição'],
        [
            ['CREATE', 'Criação de novos registros'],
            ['UPDATE', 'Atualização de registros'],
            ['DELETE', 'Exclusão de registros'],
            ['LOGIN', 'Acesso ao sistema'],
            ['LOGOUT', 'Saída do sistema'],
            ['APPROVE', 'Aprovações'],
            ['REJECT', 'Rejeições'],
            ['ARCHIVE', 'Arquivamento'],
            ['UPLOAD', 'Upload de arquivos'],
            ['EXPIRE', 'Expirações'],
            ['ANONYMIZE', 'Anonimização de dados']
        ]
    )

    add_section_title(doc, '8.3 Detalhes de Auditoria', 2)
    add_styled_paragraph(doc, 'Cada registro de auditoria contém:')
    add_bullet_list(doc, [
        'Data e hora da ação',
        'Usuário que executou a ação',
        'Tipo de papel do usuário (Admin, Recrutador, Candidato)',
        'Tipo de ação realizada',
        'Entidade afetada',
        'Endereço IP',
        'User Agent (navegador)',
        'Metadados adicionais'
    ])

    add_image_placeholder(doc, 'Figura 19: Detalhes de um registro de auditoria')

    add_page_break(doc)

    # 9. Relatórios
    add_section_title(doc, '9. RELATÓRIOS', 1)
    add_styled_paragraph(doc, 'Gere relatórios gerenciais para análise de dados da plataforma.')

    add_section_title(doc, '9.1 Acessar Relatórios', 2)
    add_numbered_list(doc, [
        'No dashboard, clique em "Relatórios Gerenciais"',
        'Selecione o tipo de relatório desejado'
    ])

    add_image_placeholder(doc, 'Figura 20: Página de relatórios')

    add_section_title(doc, '9.2 Tipos de Relatórios Disponíveis', 2)
    add_bullet_list(doc, [
        'Relatório de Vagas (por status, associado, período)',
        'Relatório de Candidaturas (por vaga, status, período)',
        'Relatório de Candidatos (por área de atuação, localização)',
        'Relatório de Recrutadores (por atividade, associado)',
        'Relatório de Associados (por atividade, vagas publicadas)'
    ])

    add_section_title(doc, '9.3 Gerar Relatório', 2)
    add_numbered_list(doc, [
        'Selecione o tipo de relatório',
        'Defina os filtros desejados (período, status, etc.)',
        'Clique em "Gerar Relatório"',
        'Visualize o relatório na tela',
        'Exporte em PDF ou Excel, se necessário'
    ])

    add_image_placeholder(doc, 'Figura 21: Geração de relatório personalizado')

    add_page_break(doc)

    # 10. FAQ
    add_section_title(doc, '10. PERGUNTAS FREQUENTES', 1)

    add_section_title(doc, 'Como adiciono um novo recrutador?', 2)
    add_styled_paragraph(doc, 'Acesse "Gerenciamento de Usuários" e clique em "Convidar Recrutador". Preencha os dados e envie o convite. O recrutador receberá um e-mail para ativar sua conta.')

    add_section_title(doc, 'O que acontece quando desativo um recrutador?', 2)
    add_styled_paragraph(doc, 'O recrutador perde acesso imediato ao sistema, mas suas vagas e candidaturas permanecem intactas. Ele pode ser reativado a qualquer momento.')

    add_section_title(doc, 'Como aprovar um pedido de associação?', 2)
    add_styled_paragraph(doc, 'Acesse "Gerenciamento de Associados" > "Pedidos Pendentes", revise os dados e clique em "Aprovar" ou "Rejeitar".')

    add_section_title(doc, 'Posso excluir uma área de atuação que tem vagas vinculadas?', 2)
    add_styled_paragraph(doc, 'Não. Primeiro é necessário desvincular todas as vagas e candidatos daquela área.')

    add_section_title(doc, 'Os logs de auditoria podem ser excluídos?', 2)
    add_styled_paragraph(doc, 'Não. Os logs são permanentes para fins de segurança e conformidade.')

    add_section_title(doc, 'Como posso exportar um relatório?', 2)
    add_styled_paragraph(doc, 'Após gerar o relatório, clique no botão "Exportar" e escolha o formato desejado (PDF ou Excel).')

    add_page_break(doc)

    # Rodapé
    add_section_title(doc, 'SUPORTE TÉCNICO', 1)
    add_styled_paragraph(doc, 'Em caso de dúvidas ou problemas técnicos, entre em contato:')
    add_bullet_list(doc, [
        'E-mail: suporte@aisam.com.br',
        'Telefone: (15) XXXX-XXXX',
        'Site: www.aisam.com.br'
    ])

    add_styled_paragraph(doc, '')
    add_styled_paragraph(doc, '© 2026 AISAM - Associação das Indústrias de Santa Maria da Serra e Região', color=CINZA_TEXTO)
    add_styled_paragraph(doc, 'Todos os direitos reservados.', color=CINZA_TEXTO)

    # Salvar documento
    doc.save('Manual_Administrador_AISAM.docx')
    print('[OK] Manual do Administrador criado com sucesso!')


def criar_manual_recrutador():
    """Cria manual do Recrutador"""
    doc = Document()

    # Configurações do documento
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Capa com logo
    logo_path = 'frontend/src/assets/aisam-logo.webp'
    add_cover_page(doc, 'MANUAL DO RECRUTADOR', 'Sistema de Gestão AISAM', logo_path)

    add_page_break(doc)

    # Página de apresentação
    add_logo_header(doc, 'MANUAL DO USUÁRIO\nRECRUTADOR')

    add_styled_paragraph(doc, 'Plataforma de Gestão de Vagas e Currículos', bold=True)
    add_styled_paragraph(doc, 'Versão 1.0 - 2026')
    add_styled_paragraph(doc, '')
    add_styled_paragraph(doc, 'Este manual destina-se aos Recrutadores que representam empresas associadas à AISAM.')

    add_page_break(doc)

    # Índice
    add_section_title(doc, 'ÍNDICE', 1)
    add_bullet_list(doc, [
        '1. Introdução',
        '2. Primeiro Acesso',
        '3. Dashboard do Recrutador',
        '4. Publicar Vagas',
        '5. Gerenciar Vagas',
        '6. Gerenciar Candidaturas',
        '7. Banco de Currículos',
        '8. Perfil do Recrutador',
        '9. Perguntas Frequentes'
    ])

    add_page_break(doc)

    # 1. Introdução
    add_section_title(doc, '1. INTRODUÇÃO', 1)
    add_styled_paragraph(doc, 'Bem-vindo à plataforma de recrutamento da AISAM!')
    add_styled_paragraph(doc, 'Como Recrutador, você pode:')
    add_bullet_list(doc, [
        'Publicar vagas de emprego para sua empresa',
        'Gerenciar o status de suas vagas (abrir, pausar, encerrar)',
        'Receber e analisar candidaturas',
        'Acessar o banco de currículos de candidatos',
        'Acompanhar o processo seletivo de cada candidato'
    ])

    add_section_title(doc, '1.1 Sobre este Manual', 2)
    add_styled_paragraph(doc, 'Este documento fornece instruções passo a passo para utilizar todas as funcionalidades disponíveis para recrutadores.')

    add_page_break(doc)

    # 2. Primeiro Acesso
    add_section_title(doc, '2. PRIMEIRO ACESSO', 1)

    add_section_title(doc, '2.1 Aceitando o Convite', 2)
    add_styled_paragraph(doc, 'Você receberá um e-mail de convite do administrador da AISAM. Para ativar sua conta:')
    add_numbered_list(doc, [
        'Abra o e-mail de convite',
        'Clique no link "Aceitar Convite"',
        'Você será redirecionado para a página de cadastro',
        'Verifique seus dados (nome e e-mail)',
        'Defina uma senha forte',
        'Confirme sua senha',
        'Clique em "Ativar Conta"'
    ])

    add_image_placeholder(doc, 'Figura 1: Tela de aceite de convite')

    add_styled_paragraph(doc, 'Atenção: O link de convite tem validade de 7 dias.', bold=True, color=AZUL_ESCURO)

    add_section_title(doc, '2.2 Fazendo Login', 2)
    add_styled_paragraph(doc, 'Após ativar sua conta:')
    add_numbered_list(doc, [
        'Acesse a URL da plataforma AISAM',
        'Clique em "Acessar Sistema"',
        'Selecione "Login Recrutador"',
        'Digite seu e-mail e senha',
        'Clique em "Entrar"'
    ])

    add_image_placeholder(doc, 'Figura 2: Tela de login do Recrutador')

    add_section_title(doc, '2.3 Esqueci Minha Senha', 2)
    add_numbered_list(doc, [
        'Na tela de login, clique em "Esqueci minha senha"',
        'Digite seu e-mail cadastrado',
        'Você receberá um e-mail com instruções',
        'Clique no link recebido',
        'Defina uma nova senha',
        'Faça login com a nova senha'
    ])

    add_image_placeholder(doc, 'Figura 3: Recuperação de senha')

    add_page_break(doc)

    # 3. Dashboard
    add_section_title(doc, '3. DASHBOARD DO RECRUTADOR', 1)
    add_styled_paragraph(doc, 'O Dashboard é sua página inicial após fazer login. Ele oferece acesso rápido às principais funcionalidades.')

    add_image_placeholder(doc, 'Figura 4: Dashboard do Recrutador', height=4.0)

    add_section_title(doc, '3.1 Cards de Acesso Rápido', 2)
    add_styled_paragraph(doc, 'O dashboard contém os seguintes cards:')
    add_bullet_list(doc, [
        'Publicar Vaga - Crie uma nova oportunidade',
        'Minhas Vagas - Visualize e gerencie suas vagas',
        'Candidaturas Recebidas - Veja os candidatos interessados',
        'Banco de Currículos - Acesse perfis de candidatos'
    ])

    add_section_title(doc, '3.2 Estatísticas Rápidas', 2)
    add_styled_paragraph(doc, 'No topo do dashboard você pode visualizar:')
    add_bullet_list(doc, [
        'Total de vagas publicadas',
        'Vagas abertas no momento',
        'Total de candidaturas recebidas',
        'Candidaturas pendentes de análise'
    ])

    add_page_break(doc)

    # 4. Publicar Vagas
    add_section_title(doc, '4. PUBLICAR VAGAS', 1)
    add_styled_paragraph(doc, 'Crie novas oportunidades de emprego para sua empresa.')

    add_section_title(doc, '4.1 Criar Nova Vaga', 2)
    add_numbered_list(doc, [
        'No dashboard, clique em "Publicar Vaga"',
        'Preencha o formulário com as informações da vaga:'
    ])

    add_styled_paragraph(doc, 'Campos do Formulário:', bold=True)
    add_bullet_list(doc, [
        'Título da Vaga (ex: Analista de Sistemas Júnior)',
        'Descrição Detalhada (responsabilidades, requisitos, benefícios)',
        'Senioridade (Estágio, Júnior, Pleno, Sênior, Especialista)',
        'Regime de Trabalho (Presencial, Híbrido, Remoto)',
        'Localidade (Cidade/Estado - opcional se remoto)',
        'E-mail para Contato',
        'Áreas de Atuação (selecione uma ou mais)',
        'Empresa Anônima (marque se não quiser exibir o nome da empresa)'
    ])

    add_image_placeholder(doc, 'Figura 5: Formulário de criação de vaga')

    add_numbered_list(doc, [
        'Revise todas as informações',
        'Clique em "Publicar Vaga"',
        'A vaga será publicada imediatamente com status ABERTA'
    ])

    add_styled_paragraph(doc, 'Dica: Seja claro e objetivo na descrição. Candidatos bem informados fazem candidaturas mais qualificadas.', bold=True, color=AZUL_CLARO)

    add_image_placeholder(doc, 'Figura 6: Confirmação de vaga publicada')

    add_page_break(doc)

    # 5. Gerenciar Vagas
    add_section_title(doc, '5. GERENCIAR VAGAS', 1)
    add_styled_paragraph(doc, 'Visualize e gerencie todas as vagas que você publicou.')

    add_section_title(doc, '5.1 Visualizar Minhas Vagas', 2)
    add_numbered_list(doc, [
        'No dashboard, clique em "Minhas Vagas"',
        'Visualize a lista de todas as suas vagas',
        'Use filtros para buscar por status ou área de atuação'
    ])

    add_image_placeholder(doc, 'Figura 7: Lista de vagas publicadas')

    add_section_title(doc, '5.2 Status das Vagas', 2)
    add_styled_paragraph(doc, 'Suas vagas podem ter os seguintes status:')

    add_table_simple(doc,
        ['Status', 'Descrição', 'Aceita Candidaturas'],
        [
            ['ABERTA', 'Vaga ativa recebendo candidaturas', 'Sim'],
            ['PAUSADA', 'Vaga temporariamente suspensa', 'Não'],
            ['FECHADA', 'Vaga encerrada definitivamente', 'Não'],
            ['ARQUIVADA', 'Vaga arquivada para referência', 'Não']
        ]
    )

    add_section_title(doc, '5.3 Editar Vaga', 2)
    add_numbered_list(doc, [
        'Na lista de vagas, clique sobre a vaga desejada',
        'Clique no botão "Editar"',
        'Atualize as informações necessárias',
        'Clique em "Salvar Alterações"'
    ])

    add_image_placeholder(doc, 'Figura 8: Edição de vaga')

    add_section_title(doc, '5.4 Pausar Vaga', 2)
    add_styled_paragraph(doc, 'Use esta opção quando precisar suspender temporariamente o recebimento de candidaturas:')
    add_numbered_list(doc, [
        'Acesse a vaga desejada',
        'Clique em "Pausar Vaga"',
        'Confirme a ação',
        'A vaga deixa de aparecer nas buscas de candidatos'
    ])

    add_styled_paragraph(doc, 'Dica: Use "Pausar" quando estiver analisando candidatos ou quando a vaga estiver temporariamente suspensa.', bold=True, color=AZUL_CLARO)

    add_image_placeholder(doc, 'Figura 9: Pausar vaga')

    add_section_title(doc, '5.5 Reabrir Vaga', 2)
    add_styled_paragraph(doc, 'Para reativar uma vaga pausada:')
    add_numbered_list(doc, [
        'Acesse a vaga pausada',
        'Clique em "Reabrir Vaga"',
        'Confirme a ação',
        'A vaga volta a aparecer nas buscas'
    ])

    add_section_title(doc, '5.6 Encerrar Vaga', 2)
    add_styled_paragraph(doc, 'Use esta opção quando a vaga for preenchida ou não for mais necessária:')
    add_numbered_list(doc, [
        'Acesse a vaga desejada',
        'Clique em "Encerrar Vaga"',
        'Confirme a ação',
        'A vaga não pode mais ser reaberta (apenas arquivada)'
    ])

    add_image_placeholder(doc, 'Figura 10: Encerrar vaga')

    add_section_title(doc, '5.7 Arquivar Vaga', 2)
    add_styled_paragraph(doc, 'Arquive vagas antigas para manter seu histórico organizado:')
    add_numbered_list(doc, [
        'Acesse a vaga encerrada',
        'Clique em "Arquivar"',
        'A vaga é movida para a seção de arquivados'
    ])

    add_section_title(doc, '5.8 Excluir Vaga', 2)
    add_styled_paragraph(doc, 'Você pode excluir vagas que criou:')
    add_numbered_list(doc, [
        'Acesse a vaga desejada',
        'Clique em "Excluir"',
        'Confirme a ação'
    ])
    add_styled_paragraph(doc, 'Atenção: Esta ação é irreversível. Todas as candidaturas associadas serão perdidas.', bold=True, color=AZUL_ESCURO)

    add_page_break(doc)

    # 6. Gerenciar Candidaturas
    add_section_title(doc, '6. GERENCIAR CANDIDATURAS', 1)
    add_styled_paragraph(doc, 'Analise e gerencie as candidaturas recebidas em suas vagas.')

    add_section_title(doc, '6.1 Visualizar Candidaturas', 2)
    add_numbered_list(doc, [
        'No dashboard, clique em "Candidaturas Recebidas"',
        'Visualize todas as candidaturas para suas vagas',
        'Use filtros para buscar por:',
        '  - Vaga específica',
        '  - Status da candidatura',
        '  - Data de candidatura'
    ])

    add_image_placeholder(doc, 'Figura 11: Lista de candidaturas recebidas')

    add_section_title(doc, '6.2 Status das Candidaturas', 2)
    add_styled_paragraph(doc, 'As candidaturas podem ter os seguintes status:')

    add_table_simple(doc,
        ['Status', 'Descrição', 'Ação do Recrutador'],
        [
            ['PENDENTE', 'Aguardando análise', 'Analisar currículo'],
            ['EM_ANALISE', 'Em processo de avaliação', 'Aprovar ou rejeitar'],
            ['ACEITA', 'Candidato aprovado', 'Entrar em contato'],
            ['RECUSADA', 'Candidato não selecionado', 'Finalizado'],
            ['CANCELADA', 'Cancelada pelo candidato', 'Finalizado']
        ]
    )

    add_section_title(doc, '6.3 Analisar Candidatura', 2)
    add_numbered_list(doc, [
        'Clique sobre a candidatura desejada',
        'Visualize o perfil completo do candidato:',
        '  - Dados pessoais (nome, e-mail, telefone)',
        '  - Cidade e estado',
        '  - Áreas de atuação',
        '  - Resumo do currículo',
        '  - Mensagem do candidato (se houver)',
        '  - Arquivo do currículo em PDF (se disponível)'
    ])

    add_image_placeholder(doc, 'Figura 12: Detalhes da candidatura')

    add_section_title(doc, '6.4 Atualizar Status da Candidatura', 2)
    add_styled_paragraph(doc, 'Para mover a candidatura no processo seletivo:')
    add_numbered_list(doc, [
        'Acesse os detalhes da candidatura',
        'Clique em "Atualizar Status"',
        'Selecione o novo status:',
        '  - EM_ANALISE: quando iniciar a avaliação',
        '  - ACEITA: quando aprovar o candidato',
        '  - RECUSADA: quando não selecionar o candidato',
        'Adicione observações (opcional)',
        'Clique em "Confirmar"'
    ])

    add_image_placeholder(doc, 'Figura 13: Atualização de status da candidatura')

    add_styled_paragraph(doc, 'Dica: Use o campo de observações para registrar feedback ou próximos passos.', bold=True, color=AZUL_CLARO)

    add_section_title(doc, '6.5 Baixar Currículo', 2)
    add_styled_paragraph(doc, 'Se o candidato fez upload do currículo em PDF:')
    add_numbered_list(doc, [
        'Acesse os detalhes da candidatura',
        'Clique no botão "Baixar Currículo"',
        'O arquivo PDF será baixado para seu computador'
    ])

    add_image_placeholder(doc, 'Figura 14: Download de currículo')

    add_page_break(doc)

    # 7. Banco de Currículos
    add_section_title(doc, '7. BANCO DE CURRÍCULOS', 1)
    add_styled_paragraph(doc, 'Acesse o banco completo de candidatos cadastrados na plataforma.')

    add_section_title(doc, '7.1 Buscar Candidatos', 2)
    add_numbered_list(doc, [
        'No dashboard, clique em "Banco de Currículos"',
        'Visualize a lista de todos os candidatos',
        'Use filtros para refinar sua busca:',
        '  - Áreas de atuação',
        '  - Localização (cidade/estado)',
        '  - Busca por palavra-chave no currículo'
    ])

    add_image_placeholder(doc, 'Figura 15: Banco de currículos')

    add_section_title(doc, '7.2 Visualizar Perfil do Candidato', 2)
    add_numbered_list(doc, [
        'Clique sobre um candidato',
        'Visualize o perfil completo',
        'Veja o histórico de candidaturas do candidato'
    ])

    add_image_placeholder(doc, 'Figura 16: Perfil do candidato')

    add_styled_paragraph(doc, 'Dica: Use o banco de currículos para buscar candidatos proativamente, mesmo antes de publicar uma vaga.', bold=True, color=AZUL_CLARO)

    add_page_break(doc)

    # 8. Perfil
    add_section_title(doc, '8. PERFIL DO RECRUTADOR', 1)

    add_section_title(doc, '8.1 Visualizar Meu Perfil', 2)
    add_numbered_list(doc, [
        'Clique no seu nome no canto superior direito',
        'Selecione "Meu Perfil"',
        'Visualize seus dados cadastrados'
    ])

    add_image_placeholder(doc, 'Figura 17: Perfil do recrutador')

    add_section_title(doc, '8.2 Editar Dados', 2)
    add_numbered_list(doc, [
        'Na página de perfil, clique em "Editar"',
        'Atualize as informações:',
        '  - Nome',
        '  - E-mail',
        'Clique em "Salvar"'
    ])

    add_section_title(doc, '8.3 Alterar Senha', 2)
    add_numbered_list(doc, [
        'Na página de perfil, clique em "Alterar Senha"',
        'Digite sua senha atual',
        'Digite a nova senha',
        'Confirme a nova senha',
        'Clique em "Salvar"'
    ])

    add_image_placeholder(doc, 'Figura 18: Alteração de senha')

    add_section_title(doc, '8.4 Sair do Sistema', 2)
    add_styled_paragraph(doc, 'Para fazer logout:')
    add_numbered_list(doc, [
        'Clique no seu nome no canto superior direito',
        'Selecione "Sair"',
        'Você será desconectado do sistema'
    ])

    add_page_break(doc)

    # 9. FAQ
    add_section_title(doc, '9. PERGUNTAS FREQUENTES', 1)

    add_section_title(doc, 'Não recebi o e-mail de convite. O que fazer?', 2)
    add_styled_paragraph(doc, 'Verifique sua caixa de spam. Se não encontrar, entre em contato com o administrador da AISAM para reenviar o convite.')

    add_section_title(doc, 'Posso publicar quantas vagas quiser?', 2)
    add_styled_paragraph(doc, 'Sim, não há limite de vagas que você pode publicar.')

    add_section_title(doc, 'Posso editar uma vaga depois de publicada?', 2)
    add_styled_paragraph(doc, 'Sim, você pode editar qualquer informação da vaga a qualquer momento.')

    add_section_title(doc, 'O que acontece se eu pausar uma vaga?', 2)
    add_styled_paragraph(doc, 'A vaga deixa de aparecer nas buscas de candidatos e não recebe mais candidaturas. Você pode reabri-la a qualquer momento.')

    add_section_title(doc, 'Qual a diferença entre pausar e encerrar?', 2)
    add_styled_paragraph(doc, 'PAUSAR é temporário e reversível. ENCERRAR é definitivo - a vaga não pode mais ser reaberta, apenas arquivada.')

    add_section_title(doc, 'Os candidatos sabem quando mudo o status da candidatura?', 2)
    add_styled_paragraph(doc, 'Sim, os candidatos podem acompanhar o status de suas candidaturas em tempo real.')

    add_section_title(doc, 'Posso ver candidatos que não se candidataram às minhas vagas?', 2)
    add_styled_paragraph(doc, 'Sim, através do Banco de Currículos você pode acessar todos os candidatos cadastrados na plataforma.')

    add_section_title(doc, 'O que é "Empresa Anônima"?', 2)
    add_styled_paragraph(doc, 'Ao marcar esta opção, o nome da sua empresa não será exibido na vaga. Aparecerá apenas "Empresa Confidencial".')

    add_page_break(doc)

    # Rodapé
    add_section_title(doc, 'SUPORTE TÉCNICO', 1)
    add_styled_paragraph(doc, 'Em caso de dúvidas ou problemas técnicos, entre em contato:')
    add_bullet_list(doc, [
        'E-mail: suporte@aisam.com.br',
        'Telefone: (15) XXXX-XXXX',
        'Site: www.aisam.com.br'
    ])

    add_styled_paragraph(doc, '')
    add_styled_paragraph(doc, '© 2026 AISAM - Associação das Indústrias de Santa Maria da Serra e Região', color=CINZA_TEXTO)
    add_styled_paragraph(doc, 'Todos os direitos reservados.', color=CINZA_TEXTO)

    # Salvar documento
    doc.save('Manual_Recrutador_AISAM.docx')
    print('[OK] Manual do Recrutador criado com sucesso!')


def criar_manual_candidato():
    """Cria manual do Candidato"""
    doc = Document()

    # Configurações do documento
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Capa com logo
    logo_path = 'frontend/src/assets/aisam-logo.webp'
    add_cover_page(doc, 'MANUAL DO CANDIDATO', 'Plataforma de Vagas AISAM', logo_path)

    add_page_break(doc)

    # Página de apresentação
    add_logo_header(doc, 'MANUAL DO USUÁRIO\nCANDIDATO')

    add_styled_paragraph(doc, 'Plataforma de Vagas de Emprego', bold=True)
    add_styled_paragraph(doc, 'Versão 1.0 - 2026')
    add_styled_paragraph(doc, '')
    add_styled_paragraph(doc, 'Este manual destina-se aos Candidatos que buscam oportunidades de emprego através da plataforma AISAM.')

    add_page_break(doc)

    # Índice
    add_section_title(doc, 'ÍNDICE', 1)
    add_bullet_list(doc, [
        '1. Introdução',
        '2. Cadastro e Primeiro Acesso',
        '3. Buscar Vagas',
        '4. Candidatar-se a Vagas',
        '5. Acompanhar Candidaturas',
        '6. Gerenciar Currículo',
        '7. Perfil do Candidato',
        '8. Perguntas Frequentes'
    ])

    add_page_break(doc)

    # 1. Introdução
    add_section_title(doc, '1. INTRODUÇÃO', 1)
    add_styled_paragraph(doc, 'Bem-vindo à plataforma de vagas da AISAM!')
    add_styled_paragraph(doc, 'Aqui você pode:')
    add_bullet_list(doc, [
        'Buscar vagas de emprego em empresas associadas à AISAM',
        'Candidatar-se às oportunidades que mais combinam com seu perfil',
        'Acompanhar o status de suas candidaturas em tempo real',
        'Manter seu currículo atualizado',
        'Receber notificações sobre suas candidaturas'
    ])

    add_section_title(doc, '1.1 Sobre a AISAM', 2)
    add_styled_paragraph(doc, 'A AISAM (Associação das Indústrias de Santa Maria da Serra e Região) conecta empresas da região com profissionais qualificados, facilitando o processo de recrutamento.')

    add_page_break(doc)

    # 2. Cadastro
    add_section_title(doc, '2. CADASTRO E PRIMEIRO ACESSO', 1)
    add_styled_paragraph(doc, 'A plataforma AISAM utiliza um sistema de acesso simplificado, sem necessidade de senha.')

    add_section_title(doc, '2.1 Criar Conta', 2)
    add_styled_paragraph(doc, 'Para se cadastrar:')
    add_numbered_list(doc, [
        'Acesse o site da AISAM',
        'Clique em "Buscar Vagas"',
        'Clique em "Cadastre-se"',
        'Preencha o formulário de cadastro:'
    ])

    add_styled_paragraph(doc, 'Campos do Formulário:', bold=True)
    add_bullet_list(doc, [
        'Nome Completo',
        'E-mail (será usado para acesso)',
        'Telefone',
        'Cidade e Estado',
        'Áreas de Atuação de Interesse (selecione uma ou mais)',
        'Resumo do Currículo (breve descrição da sua experiência)',
        'Consentimento para uso de dados (obrigatório)'
    ])

    add_image_placeholder(doc, 'Figura 1: Formulário de cadastro de candidato')

    add_numbered_list(doc, [
        'Revise as informações',
        'Clique em "Cadastrar"',
        'Você receberá um e-mail com link de acesso'
    ])

    add_section_title(doc, '2.2 Acessar o Sistema (Magic Link)', 2)
    add_styled_paragraph(doc, 'Para fazer login:')
    add_numbered_list(doc, [
        'Clique em "Acessar Sistema"',
        'Selecione "Acesso Candidato"',
        'Digite seu e-mail cadastrado',
        'Clique em "Enviar Link de Acesso"',
        'Você receberá um e-mail com o link de acesso',
        'Clique no link recebido',
        'Você será automaticamente conectado'
    ])

    add_image_placeholder(doc, 'Figura 2: Solicitação de link de acesso')

    add_styled_paragraph(doc, 'Importante: O link de acesso é válido por 24 horas. Após esse período, solicite um novo link.', bold=True, color=AZUL_ESCURO)

    add_image_placeholder(doc, 'Figura 3: E-mail com link de acesso')

    add_styled_paragraph(doc, 'Dica: Salve o link nos favoritos enquanto estiver válido para acesso rápido.', bold=True, color=AZUL_CLARO)

    add_page_break(doc)

    # 3. Buscar Vagas
    add_section_title(doc, '3. BUSCAR VAGAS', 1)
    add_styled_paragraph(doc, 'Encontre as melhores oportunidades para seu perfil.')

    add_section_title(doc, '3.1 Página de Vagas', 2)
    add_styled_paragraph(doc, 'Para buscar vagas disponíveis:')
    add_numbered_list(doc, [
        'No site principal, clique em "Vagas"',
        'Ou, após fazer login, acesse "Vagas Disponíveis"',
        'Visualize todas as vagas abertas'
    ])

    add_image_placeholder(doc, 'Figura 4: Lista de vagas disponíveis', height=4.0)

    add_section_title(doc, '3.2 Filtrar Vagas', 2)
    add_styled_paragraph(doc, 'Use os filtros para refinar sua busca:')
    add_bullet_list(doc, [
        'Regime de Trabalho: Presencial, Híbrido ou Remoto',
        'Senioridade: Estágio, Júnior, Pleno, Sênior, Especialista',
        'Áreas de Atuação: Tecnologia, Logística, Recursos Humanos, etc.',
        'Busca por Palavra-chave: Digite termos no campo de busca'
    ])

    add_image_placeholder(doc, 'Figura 5: Filtros de busca de vagas')

    add_section_title(doc, '3.3 Visualizar Detalhes da Vaga', 2)
    add_numbered_list(doc, [
        'Clique sobre uma vaga de interesse',
        'Visualize as informações completas:',
        '  - Título e descrição detalhada',
        '  - Nome da empresa (ou "Empresa Confidencial")',
        '  - Regime de trabalho',
        '  - Senioridade',
        '  - Localidade',
        '  - Áreas de atuação',
        '  - E-mail de contato',
        '  - Data de publicação'
    ])

    add_image_placeholder(doc, 'Figura 6: Detalhes completos de uma vaga')

    add_page_break(doc)

    # 4. Candidatar-se
    add_section_title(doc, '4. CANDIDATAR-SE A VAGAS', 1)
    add_styled_paragraph(doc, 'Envie sua candidatura para as vagas de seu interesse.')

    add_section_title(doc, '4.1 Fazer Candidatura', 2)
    add_styled_paragraph(doc, 'Para se candidatar a uma vaga:')
    add_numbered_list(doc, [
        'Acesse os detalhes da vaga desejada',
        'Clique no botão "Candidatar-me"',
        'Revise seus dados de perfil',
        'Adicione uma mensagem personalizada (opcional)',
        'Clique em "Confirmar Candidatura"'
    ])

    add_image_placeholder(doc, 'Figura 7: Formulário de candidatura')

    add_styled_paragraph(doc, 'Dica: Use a mensagem personalizada para destacar por que você é o candidato ideal para a vaga.', bold=True, color=AZUL_CLARO)

    add_section_title(doc, '4.2 Confirmação', 2)
    add_styled_paragraph(doc, 'Após confirmar:')
    add_bullet_list(doc, [
        'Você verá uma mensagem de confirmação',
        'Sua candidatura será enviada ao recrutador',
        'Você pode acompanhar o status em "Minhas Candidaturas"'
    ])

    add_image_placeholder(doc, 'Figura 8: Confirmação de candidatura enviada')

    add_section_title(doc, '4.3 Limitações', 2)
    add_styled_paragraph(doc, 'Importante saber:')
    add_bullet_list(doc, [
        'Você só pode se candidatar uma vez para cada vaga',
        'Não é possível editar a candidatura após o envio',
        'Você pode cancelar a candidatura enquanto estiver pendente'
    ])

    add_page_break(doc)

    # 5. Acompanhar Candidaturas
    add_section_title(doc, '5. ACOMPANHAR CANDIDATURAS', 1)
    add_styled_paragraph(doc, 'Monitore o status de todas as suas candidaturas em tempo real.')

    add_section_title(doc, '5.1 Visualizar Minhas Candidaturas', 2)
    add_numbered_list(doc, [
        'Faça login no sistema',
        'Clique em "Minhas Candidaturas"',
        'Visualize todas as vagas para as quais se candidatou'
    ])

    add_image_placeholder(doc, 'Figura 9: Lista de candidaturas do candidato')

    add_section_title(doc, '5.2 Status das Candidaturas', 2)
    add_styled_paragraph(doc, 'Suas candidaturas podem ter os seguintes status:')

    add_table_simple(doc,
        ['Status', 'Significado', 'O que fazer'],
        [
            ['PENDENTE', 'Aguardando análise do recrutador', 'Aguarde'],
            ['EM_ANALISE', 'Recrutador está avaliando seu perfil', 'Aguarde'],
            ['ACEITA', 'Você foi aprovado!', 'Aguarde contato'],
            ['RECUSADA', 'Não foi selecionado desta vez', 'Continue buscando'],
            ['CANCELADA', 'Você cancelou a candidatura', 'Finalizado']
        ]
    )

    add_section_title(doc, '5.3 Detalhes da Candidatura', 2)
    add_numbered_list(doc, [
        'Clique sobre uma candidatura',
        'Visualize:',
        '  - Detalhes completos da vaga',
        '  - Status atual da candidatura',
        '  - Data da candidatura',
        '  - Observações do recrutador (se houver)',
        '  - Sua mensagem enviada'
    ])

    add_image_placeholder(doc, 'Figura 10: Detalhes da candidatura')

    add_section_title(doc, '5.4 Cancelar Candidatura', 2)
    add_styled_paragraph(doc, 'Você pode cancelar uma candidatura enquanto ela estiver com status PENDENTE ou EM_ANALISE:')
    add_numbered_list(doc, [
        'Acesse os detalhes da candidatura',
        'Clique em "Cancelar Candidatura"',
        'Confirme a ação',
        'O status mudará para CANCELADA'
    ])

    add_image_placeholder(doc, 'Figura 11: Cancelamento de candidatura')

    add_styled_paragraph(doc, 'Atenção: Não é possível desfazer o cancelamento. Você precisará se candidatar novamente.', bold=True, color=AZUL_ESCURO)

    add_page_break(doc)

    # 6. Gerenciar Currículo
    add_section_title(doc, '6. GERENCIAR CURRÍCULO', 1)
    add_styled_paragraph(doc, 'Mantenha seu perfil e currículo sempre atualizados.')

    add_section_title(doc, '6.1 Acessar Meu Currículo', 2)
    add_numbered_list(doc, [
        'Faça login no sistema',
        'Clique em "Meu Currículo"',
        'Visualize seu perfil completo'
    ])

    add_image_placeholder(doc, 'Figura 12: Página do currículo')

    add_section_title(doc, '6.2 Editar Informações', 2)
    add_styled_paragraph(doc, 'Para atualizar seus dados:')
    add_numbered_list(doc, [
        'Na página "Meu Currículo", clique em "Editar"',
        'Atualize as informações:',
        '  - Nome completo',
        '  - Telefone',
        '  - Cidade e estado',
        '  - Áreas de atuação',
        '  - Resumo do currículo',
        'Clique em "Salvar"'
    ])

    add_image_placeholder(doc, 'Figura 13: Edição de dados do currículo')

    add_styled_paragraph(doc, 'Dica: Mantenha seu resumo atualizado com suas experiências mais recentes.', bold=True, color=AZUL_CLARO)

    add_section_title(doc, '6.3 Upload de Currículo em PDF', 2)
    add_styled_paragraph(doc, 'Você pode fazer upload do seu currículo completo em formato PDF:')
    add_numbered_list(doc, [
        'Na página "Meu Currículo", clique em "Upload de Currículo"',
        'Selecione um arquivo PDF do seu computador',
        'Aguarde o upload completar',
        'Clique em "Salvar"'
    ])

    add_image_placeholder(doc, 'Figura 14: Upload de currículo em PDF')

    add_styled_paragraph(doc, 'Importante: O arquivo deve estar em formato PDF e ter no máximo 5MB.', bold=True, color=AZUL_ESCURO)

    add_section_title(doc, '6.4 Visualizar Currículo', 2)
    add_styled_paragraph(doc, 'Veja como seu perfil aparece para os recrutadores:')
    add_numbered_list(doc, [
        'Acesse "Meu Currículo"',
        'Clique em "Visualizar como Recrutador"',
        'Veja exatamente o que os recrutadores enxergam'
    ])

    add_image_placeholder(doc, 'Figura 15: Visualização do perfil do candidato')

    add_page_break(doc)

    # 7. Perfil
    add_section_title(doc, '7. PERFIL DO CANDIDATO', 1)

    add_section_title(doc, '7.1 Alterar E-mail', 2)
    add_styled_paragraph(doc, 'Para alterar seu e-mail de acesso:')
    add_numbered_list(doc, [
        'Acesse "Meu Perfil"',
        'Clique em "Alterar E-mail"',
        'Digite o novo e-mail',
        'Confirme o e-mail',
        'Clique em "Salvar"',
        'Você receberá um e-mail de confirmação'
    ])

    add_image_placeholder(doc, 'Figura 16: Alteração de e-mail')

    add_section_title(doc, '7.2 Consentimento de Dados', 2)
    add_styled_paragraph(doc, 'Seus dados são protegidos de acordo com a LGPD (Lei Geral de Proteção de Dados):')
    add_bullet_list(doc, [
        'Seus dados pessoais são usados apenas para fins de recrutamento',
        'Apenas recrutadores autorizados podem visualizar seu perfil',
        'Você pode solicitar exclusão de seus dados a qualquer momento',
        'Seus dados ficam ativos por 30 dias após o último acesso'
    ])

    add_section_title(doc, '7.3 Excluir Conta', 2)
    add_styled_paragraph(doc, 'Para excluir sua conta permanentemente:')
    add_numbered_list(doc, [
        'Acesse "Meu Perfil"',
        'Clique em "Excluir Conta"',
        'Leia o aviso sobre exclusão de dados',
        'Confirme a exclusão',
        'Todas as suas candidaturas serão canceladas',
        'Seus dados serão removidos do sistema'
    ])

    add_styled_paragraph(doc, 'Atenção: Esta ação é irreversível e não pode ser desfeita.', bold=True, color=AZUL_ESCURO)

    add_section_title(doc, '7.4 Sair do Sistema', 2)
    add_styled_paragraph(doc, 'Para fazer logout:')
    add_numbered_list(doc, [
        'Clique no seu nome no canto superior direito',
        'Selecione "Sair"',
        'Você será desconectado do sistema'
    ])

    add_page_break(doc)

    # 8. FAQ
    add_section_title(doc, '8. PERGUNTAS FREQUENTES', 1)

    add_section_title(doc, 'Por que não preciso de senha?', 2)
    add_styled_paragraph(doc, 'A plataforma usa um sistema de "Magic Link" que envia um link único de acesso para seu e-mail, tornando o processo mais seguro e simples.')

    add_section_title(doc, 'O link de acesso expirou. O que fazer?', 2)
    add_styled_paragraph(doc, 'Solicite um novo link na página de login. Você pode fazer isso quantas vezes precisar.')

    add_section_title(doc, 'Posso me candidatar a quantas vagas quiser?', 2)
    add_styled_paragraph(doc, 'Sim, não há limite de candidaturas. Candidate-se a todas as vagas que correspondam ao seu perfil.')

    add_section_title(doc, 'Posso editar minha candidatura depois de enviada?', 2)
    add_styled_paragraph(doc, 'Não, mas você pode cancelá-la e fazer uma nova candidatura se a vaga ainda estiver aberta.')

    add_section_title(doc, 'Como sei se o recrutador viu minha candidatura?', 2)
    add_styled_paragraph(doc, 'Quando o recrutador mudar o status da candidatura para "EM_ANALISE", você saberá que ele está avaliando seu perfil.')

    add_section_title(doc, 'O que significa "Empresa Confidencial"?', 2)
    add_styled_paragraph(doc, 'Algumas empresas optam por não divulgar seu nome na vaga. Você saberá a identidade da empresa caso seja aprovado.')

    add_section_title(doc, 'Preciso fazer upload do currículo em PDF?', 2)
    add_styled_paragraph(doc, 'Não é obrigatório, mas é altamente recomendado. Isso dá aos recrutadores uma visão mais completa do seu perfil.')

    add_section_title(doc, 'Meus dados ficam visíveis para todos?', 2)
    add_styled_paragraph(doc, 'Não. Apenas recrutadores autenticados e autorizados podem visualizar seu perfil.')

    add_section_title(doc, 'Por quanto tempo minha conta fica ativa?', 2)
    add_styled_paragraph(doc, 'Sua conta permanece ativa por 30 dias desde o último acesso. Após isso, você precisará solicitar um novo link para reativar.')

    add_section_title(doc, 'Recebo notificações sobre minhas candidaturas?', 2)
    add_styled_paragraph(doc, 'Sim, você pode receber e-mails quando o status de suas candidaturas for atualizado.')

    add_page_break(doc)

    # Rodapé
    add_section_title(doc, 'SUPORTE TÉCNICO', 1)
    add_styled_paragraph(doc, 'Em caso de dúvidas ou problemas técnicos, entre em contato:')
    add_bullet_list(doc, [
        'E-mail: suporte@aisam.com.br',
        'Telefone: (15) XXXX-XXXX',
        'Site: www.aisam.com.br'
    ])

    add_styled_paragraph(doc, '')
    add_styled_paragraph(doc, 'BOA SORTE NA SUA BUSCA POR OPORTUNIDADES!', bold=True, color=AZUL_PRINCIPAL)
    add_styled_paragraph(doc, '')
    add_styled_paragraph(doc, '© 2026 AISAM - Associação das Indústrias de Santa Maria da Serra e Região', color=CINZA_TEXTO)
    add_styled_paragraph(doc, 'Todos os direitos reservados.', color=CINZA_TEXTO)

    # Salvar documento
    doc.save('Manual_Candidato_AISAM.docx')
    print('[OK] Manual do Candidato criado com sucesso!')


if __name__ == '__main__':
    print('Gerando Manuais de Usuário da Plataforma AISAM...\n')

    criar_manual_admin()
    criar_manual_recrutador()
    criar_manual_candidato()

    print('\n[SUCESSO] Todos os manuais foram gerados com sucesso!')
    print('\nArquivos criados:')
    print('  1. Manual_Administrador_AISAM.docx')
    print('  2. Manual_Recrutador_AISAM.docx')
    print('  3. Manual_Candidato_AISAM.docx')
    print('\nOs manuais incluem:')
    print('  [OK] Cores da marca AISAM')
    print('  [OK] Espacos reservados para insercao de prints de tela')
    print('  [OK] Instrucoes detalhadas para cada funcionalidade')
    print('  [OK] Secoes de FAQ e suporte')
